import os
import tempfile
from typing import Tuple
import httpx
import re
import PyPDF2
import docx
import io


async def download_file(url: str) -> str:
    """
    Download a file from a URL and save it to a temporary file.
    
    Args:
        url: URL of the file to download
        
    Returns:
        Path to the temporary file
    """
    async with httpx.AsyncClient() as client:
        response = await client.get(url)
        response.raise_for_status()
        
        # Create a temporary file
        fd, path = tempfile.mkstemp()
        try:
            with os.fdopen(fd, 'wb') as tmp:
                tmp.write(response.content)
            return path
        except Exception as e:
            os.unlink(path)
            raise e


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text from the PDF
    """
    text = ""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Extract text from each page
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
                
        return text
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return f"Error extracting text from PDF: {e}"


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text from the DOCX
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
                
        return text
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        return f"Error extracting text from DOCX: {e}"


async def extract_text_from_resume(file_path: str) -> str:
    """
    Extract text from a resume file.
    
    This implementation handles different file types:
    - PDF files using PyPDF2
    - DOCX files using python-docx
    - Text files directly
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text from the resume
    """
    try:
        # Try to determine the file type based on extension
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # For PDF files
        if file_extension == '.pdf':
            return extract_text_from_pdf(file_path)
            
        # For DOCX files
        elif file_extension == '.docx':
            return extract_text_from_docx(file_path)
            
        # For text files, read directly
        elif file_extension in ['.txt', '.md', '.csv', '.json']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
        
        # For other file types, try to read as text
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
            
    except UnicodeDecodeError:
        # If we can't read it as text, try to read it as binary
        try:
            with open(file_path, 'rb') as file:
                # Try to detect if it's a PDF by checking the magic number
                content = file.read(4)
                file.seek(0)
                
                if content == b'%PDF':
                    # It's a PDF file with wrong extension
                    return extract_text_from_pdf(file_path)
                
                # For other binary files, return a message
                return "This file type is not supported for text extraction. Please upload a PDF, DOCX, or text file."
        except Exception as e:
            print(f"Error extracting text from file: {e}")
            return f"Error extracting text from file: {e}"
    except Exception as e:
        print(f"Error extracting text from resume: {e}")
        return f"Error extracting text from resume: {e}"


async def process_resume(file_path: str) -> Tuple[str, str]:
    """
    Process a resume file by extracting text.
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Tuple of (file_path, extracted_text)
    """
    try:
        # Extract text from the file
        text = await extract_text_from_resume(file_path)
        
        # For debugging
        print(f"Extracted text from resume: {text[:100]}...")
        
        return file_path, text
    except Exception as e:
        print(f"Error processing resume: {e}")
        return file_path, f"Error processing resume: {e}" 